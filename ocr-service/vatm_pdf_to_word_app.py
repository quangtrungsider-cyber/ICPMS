"""
App giao diện: Bóc tách PDF tiếng Việt -> Word
- Chọn file PDF, chọn trang (tất cả hoặc một số trang cụ thể)
- Dùng PP-StructureV3 để lấy cấu trúc/thứ tự đọc, VietOCR để nhận diện chữ
- (Tuỳ chọn) Dùng Claude API để chuẩn hoá heading & sửa lỗi OCR còn sót trước khi ghi ra Word

CÀI ĐẶT TRƯỚC KHI CHẠY:
    pip install paddlepaddle
    pip install "paddleocr[doc-parser]"
    pip install pdf2image python-docx beautifulsoup4 opencv-python pillow numpy pypdf
    pip install torch
    pip install vietocr
    pip install anthropic        (chỉ cần nếu dùng tuỳ chọn AI)

Cần cài Poppler (https://github.com/oschwartz10612/poppler-windows) và sửa đúng
POPPLER_PATH dưới đây.
"""

import os
import sys

# PHẢI đặt TRƯỚC khi import numpy/torch/paddle/opencv: paddlepaddle và torch đều
# tự mang theo bản riêng của thư viện OpenMP (libiomp5/MKL). Khi cả hai cùng nạp
# vào một process trên Windows, chúng có thể đụng nhau và làm crash thẳng cả
# chương trình (không kịp ném lỗi Python, app tự tắt im lặng). Biến môi trường
# này cho phép nạp trùng mà không crash.
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"


def _bundled_dir_candidates(folder_name):
    """Trả về danh sách các vị trí CÓ THỂ chứa 1 thư mục đóng gói kèm (poppler/
    models/paddlex_cache), theo thứ tự ưu tiên. PyInstaller bản mới (>=6) gói
    phần lớn file vào thư mục con "_internal" CẠNH file .exe, không phải ngay
    cạnh .exe như bản cũ - nên phải kiểm tra CẢ 2 chỗ, không chỉ 1."""
    base_dir = os.path.dirname(sys.executable if getattr(sys, "frozen", False) else os.path.abspath(__file__))
    return [
        os.path.join(base_dir, folder_name),
        os.path.join(base_dir, "_internal", folder_name),
    ]


def _find_bundled_dir(folder_name):
    for candidate in _bundled_dir_candidates(folder_name):
        if os.path.isdir(candidate):
            return candidate
    return None


# PHẢI đặt TRƯỚC khi import paddleocr/paddlex: paddlex đọc biến môi trường này
# NGAY LÚC import để quyết định lưu/tìm model AI (PP-StructureV3...) ở đâu (mặc
# định là C:\Users\<tên>\.paddlex, tự tải về ~1.9GB từ Internet ở lần chạy đầu
# trên MỖI máy). Nếu anh đã đóng gói sẵn thư mục "paddlex_cache" (chứa sẵn
# official_models đã tải) kèm theo app, paddlex sẽ tự dùng luôn, không tải lại -
# máy nào mở app lên cũng chạy được ngay, không cần Internet.
_bundled_paddlex_cache = _find_bundled_dir("paddlex_cache")
if _bundled_paddlex_cache:
    os.environ["PADDLE_PDX_CACHE_HOME"] = _bundled_paddlex_cache

import re
import json
import queue
import threading
import traceback

import cv2
import numpy as np
from PIL import Image
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pdf2image import convert_from_path
from pypdf import PdfReader

import tkinter as tk
from tkinter import ttk, filedialog, messagebox


# ----------------------------------------------------------------------------
# CẤU HÌNH
# ----------------------------------------------------------------------------
def _find_poppler_path():
    """Tìm Poppler theo thứ tự ưu tiên:
    1) Thư mục "poppler\\Library\\bin" đóng gói kèm app (cạnh .exe hoặc trong
       _internal, tuỳ bản PyInstaller) - để máy khác mở lên dùng được ngay,
       không cần ai tự cài Poppler riêng.
    2) Đường dẫn cố định C:\\poppler\\Library\\bin (máy cài thử nghiệm đầu tiên).
    3) Trả về None - khi đó pdf2image sẽ tự tìm Poppler trong PATH của hệ điều
       hành (vẫn chạy được nếu máy đó đã cài Poppler và thêm vào PATH)."""
    for base in _bundled_dir_candidates("poppler"):
        candidate = os.path.join(base, "Library", "bin")
        if os.path.isdir(candidate):
            return candidate
    if os.path.isdir(r"C:\poppler\Library\bin"):
        return r"C:\poppler\Library\bin"
    return None


POPPLER_PATH = _find_poppler_path()   # None nghĩa là để pdf2image tự tìm trong PATH hệ thống

TITLE_LABELS = {"doc_title", "paragraph_title", "abstract_title", "reference_title", "content_title"}
TEXT_LABELS = {"text"}
TABLE_LABELS = {"table"}
# header/header_image (quốc hiệu, tên cơ quan, ngày tháng...) là NỘI DUNG THẬT,
# vẫn cần xuất hiện trong văn bản -> không loại bỏ, chỉ coi như đoạn văn thường.
# Chỉ loại các loại thực sự là rác lặp lại không cần giữ: footer, chú thích chân
# trang, dấu mộc/con dấu, số trang.
HEADER_AS_TEXT_LABELS = {"header", "header_image"}
DROP_LABELS = {"footer", "footer_image", "footnote", "aside_text", "seal", "number"}

FONT = "Times New Roman"   # font chuẩn văn bản hành chính Việt Nam


# ----------------------------------------------------------------------------
# DỰNG WORD THEO ĐÚNG KHUÔN VĂN BẢN HÀNH CHÍNH VIỆT NAM
# (khối cơ quan/quốc hiệu 2 cột, tiêu đề căn giữa, "Kính gửi" căn giữa,
#  thân bài thụt đầu dòng + canh đều hai bên)
# ----------------------------------------------------------------------------
QUOCHIEU_TEXT = "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM"   # chuỗi cố định, không bao giờ thay đổi
DATE_PATTERN = re.compile(r"ngày\s+\d{1,2}\s+tháng\s+\d{1,2}\s+năm\s+\d{4}", re.IGNORECASE)
TIEUNGU_PATTERN = re.compile(r"(Độc lập\s*[-–—]\s*Tự do\s*[-–—]\s*Hạnh phúc)", re.IGNORECASE)
SOHIEU_PATTERN = re.compile(r"số\s*:?\s*[\w]+[/-][\w-]+", re.IGNORECASE)
KINHGUI_PATTERN = re.compile(r"^\s*kính\s+gửi", re.IGNORECASE)
NOINHAN_PATTERN = re.compile(r"^\s*nơi\s+nhận\s*:?", re.IGNORECASE)
# CHỮ HOA toàn bộ, không bắt re.IGNORECASE: dòng chức danh người ký luôn viết hoa
# (TM./KT. BAN..., GIÁM ĐỐC, TRƯỞNG BAN...), còn nhắc chức danh trong câu thường
# ("...kính đề nghị Tổng giám đốc...") thì viết hoa-thường bình thường -> không khớp,
# tránh ghép nhầm như "Tổng giám đốc" giữa câu vào khối ký tên.
SIGNATURE_ROLE_PATTERN = re.compile(
    r"\b(TM\.|KT\.|Q\.|GIÁM ĐỐC|TỔNG GIÁM ĐỐC|TRƯỞNG BAN|PHÓ TRƯỞNG BAN|CHỦ TỊCH|PHÓ CHỦ TỊCH)\b"
)
# tách dòng "TỜ TRÌNH"/"CÔNG VĂN"/... (loại văn bản, viết hoa, ngắn) ra khỏi phần
# trích yếu theo sau (thường bắt đầu bằng "Về việc" hoặc "V/v")
DOC_TYPE_SPLIT_PATTERN = re.compile(r"^([A-ZÀ-Ỹ0-9\s]{3,40}?)\s+(Về việc|V/v\.?)\s*(.*)$", re.UNICODE)


def _set_run_font(run, size=13, bold=False, italic=False, underline=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT)   # ép font cả phần chữ có dấu, tránh Word tự đổi font khác


def _add_para(doc_or_cell, text="", align=WD_ALIGN_PARAGRAPH.LEFT, size=13, bold=False,
              italic=False, underline=False, space_after=8, space_before=0,
              first_line_indent=None, line_spacing=1.3, reuse_first=False):
    if reuse_first and len(doc_or_cell.paragraphs) == 1 and not doc_or_cell.paragraphs[0].runs:
        p = doc_or_cell.paragraphs[0]
    else:
        p = doc_or_cell.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)
    if text:
        run = p.add_run(text)
        _set_run_font(run, size=size, bold=bold, italic=italic, underline=underline)
    return p


def _remove_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)


def _add_letterhead(doc, item):
    """Khối đầu trang 2 cột: cơ quan ban hành (trái) / quốc hiệu - tiêu ngữ - ngày (phải)."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(6)
    table.columns[1].width = Cm(10)
    _remove_table_borders(table)

    left = table.rows[0].cells[0]
    left.width = Cm(6)
    first = True
    for line in (item.get("org_lines") or []):
        _add_para(left, line, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0,
                  line_spacing=1.15, reuse_first=first)
        first = False
    if item.get("so_hieu"):
        _add_para(left, item["so_hieu"], align=WD_ALIGN_PARAGRAPH.CENTER, italic=True,
                  space_after=0, line_spacing=1.15, reuse_first=first)

    right = table.rows[0].cells[1]
    right.width = Cm(10)
    _add_para(right, item.get("quoc_hieu") or QUOCHIEU_TEXT, align=WD_ALIGN_PARAGRAPH.CENTER,
              bold=True, space_after=0, line_spacing=1.15, reuse_first=True)
    if item.get("tieu_ngu"):
        _add_para(right, item["tieu_ngu"], align=WD_ALIGN_PARAGRAPH.CENTER, bold=True,
                  underline=True, space_after=0, line_spacing=1.15)
    if item.get("ngay_thang"):
        _add_para(right, "")
        _add_para(right, item["ngay_thang"], align=WD_ALIGN_PARAGRAPH.CENTER, italic=True,
                  space_after=0, line_spacing=1.15)

    _add_para(doc, "", space_after=6)


def _add_signature_block(doc, item):
    """Khối cuối văn bản 2 cột: 'Nơi nhận:' + danh sách (trái) / chức danh người
    ký + tên (phải, căn giữa, in đậm)."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(8)
    table.columns[1].width = Cm(8)
    _remove_table_borders(table)

    left = table.rows[0].cells[0]
    left.width = Cm(8)
    if item.get("noi_nhan_items") is not None:
        items = [p.strip(" ;.") for p in item["noi_nhan_items"] if p and p.strip(" ;.")]
    else:
        rest = NOINHAN_PATTERN.sub("", item.get("noi_nhan_text", "")).strip()
        items = [p.strip(" ;.") for p in re.split(r"\s*-\s+", rest) if p.strip(" ;.")]
    _add_para(left, "Nơi nhận:", align=WD_ALIGN_PARAGRAPH.LEFT, italic=True, size=12,
              space_after=2, line_spacing=1.1, reuse_first=True)
    for line in items:
        _add_para(left, f"- {line};", align=WD_ALIGN_PARAGRAPH.LEFT, italic=True, size=12,
                  space_after=0, line_spacing=1.1)

    right = table.rows[0].cells[1]
    right.width = Cm(8)
    sig_text = item.get("signature_text", "")
    if sig_text:
        first = True
        for line in sig_text.split("\n"):
            if line.strip():
                _add_para(right, line.strip(), align=WD_ALIGN_PARAGRAPH.CENTER, bold=True,
                          size=13, space_after=0, line_spacing=1.3, reuse_first=first)
            else:
                _add_para(right, "", space_after=0, reuse_first=first)
            first = False
    else:
        right.paragraphs[0].text = ""


def restructure_document(render_items):
    """Quét toàn bộ các khối heading/paragraph đã trích xuất (bất kể đang ở vị trí
    nào trong danh sách - vì bước vớt dòng mồ côi có thể đẩy chúng lệch chỗ), tìm
    đúng 2 khối: tên cơ quan + số hiệu văn bản, và quốc hiệu + tiêu ngữ + ngày
    tháng (nhận diện bằng chuỗi quốc hiệu CỐ ĐỊNH, không bao giờ sai), gộp lại
    thành 1 khối "letterhead" 2 cột chèn ở đầu văn bản. Nếu không tìm thấy quốc
    hiệu thì coi như văn bản không theo khuôn này -> giữ nguyên, không ép."""
    n = len(render_items)

    def is_text(i):
        return render_items[i]["kind"] in ("heading", "paragraph")

    quochieu_idx = next((i for i in range(n) if is_text(i) and
                          QUOCHIEU_TEXT in render_items[i]["text"].upper()), None)
    if quochieu_idx is None:
        return render_items   # không thấy quốc hiệu -> không đủ tin cậy để dựng lại, giữ nguyên

    quochieu_text = render_items[quochieu_idx]["text"]
    tieu_ngu = ""
    m = TIEUNGU_PATTERN.search(quochieu_text)
    if m:
        tieu_ngu = re.sub(r"\s+", " ", m.group(1)).strip()
        quochieu_text = quochieu_text[:m.start()].strip()

    ngay_thang = ""
    date_idx = None
    m2 = DATE_PATTERN.search(render_items[quochieu_idx]["text"])
    if m2:
        ngay_thang = m2.group(0).strip()
    else:
        search_order = (list(range(max(0, quochieu_idx - 3), quochieu_idx)) +
                         list(range(quochieu_idx + 1, min(n, quochieu_idx + 4))))
        for i in search_order:
            if is_text(i) and DATE_PATTERN.search(render_items[i]["text"]):
                date_idx = i
                ngay_thang = render_items[i]["text"].strip()   # giữ nguyên cả "Hà Nội, ngày ..."
                break

    title_idx = next((i for i in range(n) if render_items[i]["kind"] == "heading"), None)
    search_limit = quochieu_idx if title_idx is None else min(quochieu_idx, title_idx)
    org_idx = None
    org_lines = []
    so_hieu = ""
    for i in range(search_limit):
        if not is_text(i) or len(render_items[i]["text"]) > 200:
            continue
        m4 = SOHIEU_PATTERN.search(render_items[i]["text"])
        if m4:
            org_idx = i
            so_hieu = re.sub(r"^số", "Số", render_items[i]["text"][m4.start():].strip(),
                              flags=re.IGNORECASE)
            org_part = render_items[i]["text"][:m4.start()].strip()
            org_lines = [org_part] if org_part else []
            break

    letterhead_item = {
        "kind": "letterhead", "org_lines": org_lines, "so_hieu": so_hieu,
        "quoc_hieu": quochieu_text, "tieu_ngu": tieu_ngu, "ngay_thang": ngay_thang,
        "page": render_items[org_idx].get("page") if org_idx is not None else render_items[quochieu_idx].get("page"),
    }

    consumed = {quochieu_idx}
    if date_idx is not None:
        consumed.add(date_idx)
    if org_idx is not None:
        consumed.add(org_idx)

    new_items = [letterhead_item] + [render_items[i] for i in range(n) if i not in consumed]
    return new_items


def _looks_like_signer_name(text):
    """Dòng tên người ký thường tách riêng khỏi dòng chức danh (cách nhau bởi
    chữ ký/khoảng trống), ví dụ 'Đỗ Thuận An' - 2-5 từ, mỗi từ viết hoa đầu,
    không số, không dấu câu kiểu liệt kê."""
    text = text.strip()
    if not text or len(text) > 60:
        return False
    if any(ch.isdigit() for ch in text):
        return False
    if any(p in text for p in (";", ":", "/", ",")):
        return False
    words = text.split()
    if not (2 <= len(words) <= 5):
        return False
    return all(w[0].isupper() for w in words if w)


def restructure_signature_blocks(render_items):
    """Tìm khối 'Nơi nhận:' (chuỗi rất đặc trưng, hầu như chỉ xuất hiện ở cuối văn
    bản hành chính) và khối chức danh người ký gần đó (TM./KT./GIÁM ĐỐC/...),
    gộp lại thành 1 bảng 2 cột giống bản gốc. Có thể có nhiều cặp trong 1 file
    (văn bản chính + các phụ lục kèm theo, mỗi phần có chữ ký riêng). Tên người
    ký thường nằm tách riêng (cách dòng chức danh bởi chữ ký/khoảng trống) nên
    tìm thêm và gộp vào cùng khối, tránh bị lạc thành đoạn văn riêng."""
    n = len(render_items)

    def is_text(i):
        return render_items[i]["kind"] in ("heading", "paragraph")

    noinhan_indices = [i for i in range(n) if is_text(i) and NOINHAN_PATTERN.match(render_items[i]["text"])]
    if not noinhan_indices:
        return render_items

    pairs = []
    used = set()
    for ni in noinhan_indices:
        window = list(range(max(0, ni - 3), ni)) + list(range(ni + 1, min(n, ni + 4)))
        sig_idx = None
        for j in window:
            if j in used or j in noinhan_indices or not is_text(j):
                continue
            if len(render_items[j]["text"]) < 250 and SIGNATURE_ROLE_PATTERN.search(render_items[j]["text"]):
                sig_idx = j
                break

        name_idx = None
        if sig_idx is not None:
            name_window = (list(range(sig_idx + 1, min(n, sig_idx + 4))) +
                            list(range(max(0, sig_idx - 3), sig_idx)))
            for j in name_window:
                if j in used or j in noinhan_indices or not is_text(j):
                    continue
                if _looks_like_signer_name(render_items[j]["text"]):
                    name_idx = j
                    break

        pairs.append((ni, sig_idx, name_idx))
        used.add(ni)
        if sig_idx is not None:
            used.add(sig_idx)
        if name_idx is not None:
            used.add(name_idx)

    pairs_by_first = {ni: (sig, name) for ni, sig, name in pairs}
    consumed_extra = ({sig for _, sig, _ in pairs if sig is not None} |
                       {name for _, _, name in pairs if name is not None})

    new_items = []
    for i in range(n):
        if i in consumed_extra:
            continue   # đã/sẽ được gộp chung với "Nơi nhận:" tương ứng
        if i in pairs_by_first:
            sig_idx, name_idx = pairs_by_first[i]
            sig_text = render_items[sig_idx]["text"] if sig_idx is not None else ""
            if name_idx is not None:
                name_text = render_items[name_idx]["text"].strip()
                sig_text = (sig_text + "\n\n" + name_text) if sig_text else name_text
            new_items.append({
                "kind": "signature_block",
                "noi_nhan_text": render_items[i]["text"],
                "signature_text": sig_text,
                "page": render_items[i].get("page"),
            })
            continue
        new_items.append(render_items[i])
    return new_items


def render_docx(render_items, output_path, page_orientation=None):
    page_orientation = page_orientation or {}
    doc = Document()

    def setup_portrait(section):
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    def setup_landscape(section):
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    setup_portrait(doc.sections[0])
    current_is_landscape = False

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(13)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT)

    # Heading ĐẦU TIÊN trong văn bản = tên loại văn bản chính (TỜ TRÌNH/CÔNG VĂN/...)
    # -> căn giữa, in đậm. Các heading sau đó là tiêu đề mục trong thân bài
    # (1., 2., II., IV. ...) -> căn trái, in đậm, không tách dòng trích yếu.
    first_heading_idx = next((i for i, it in enumerate(render_items) if it["kind"] == "heading"), None)

    for idx, item in enumerate(render_items):
        kind = item["kind"]

        wants_landscape = bool(page_orientation.get(item.get("page"), False))
        if wants_landscape != current_is_landscape:
            new_section = doc.add_section(WD_SECTION.NEW_PAGE)
            if wants_landscape:
                setup_landscape(new_section)
            else:
                setup_portrait(new_section)
            current_is_landscape = wants_landscape

        if kind == "letterhead":
            _add_letterhead(doc, item)

        elif kind == "signature_block":
            _add_signature_block(doc, item)

        elif kind == "table":
            rows = item["rows"]
            if not rows:
                continue
            cell_size = item.get("font_size") or 12
            ncols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=ncols)
            table.style = "Table Grid"

            # Độ rộng cột theo đúng tỉ lệ đo được từ PDF gốc (nếu xác định được),
            # quy đổi theo đúng phần chiều rộng còn lại của trang (đã trừ lề) -
            # để cột hẹp (STT) vẫn hẹp, cột rộng (nội dung) vẫn rộng như bản gốc.
            ratios = item.get("col_width_ratios")
            if ratios and len(ratios) == ncols:
                avail_cm = 26.7 if current_is_landscape else 16.0   # = page_width - left - right margin
                col_widths_cm = [max(1.0, r * avail_cm) for r in ratios]
                table.autofit = False
                table.allow_autofit = False
                for c_idx, w_cm in enumerate(col_widths_cm):
                    table.columns[c_idx].width = Cm(w_cm)

            for r_idx, row in enumerate(rows):
                for c_idx, cell_text in enumerate(row):
                    cell = table.cell(r_idx, c_idx)
                    cell.text = ""
                    if ratios and len(ratios) == ncols:
                        cell.width = Cm(col_widths_cm[c_idx])   # python-docx cần set cả ở cell mới ăn chắc
                    _add_para(cell, cell_text, align=WD_ALIGN_PARAGRAPH.LEFT, size=cell_size,
                              space_after=0, line_spacing=1.1, reuse_first=True)

        elif kind == "heading":
            text = item["text"]
            is_title = item.get("is_title")
            if is_title is None:
                is_title = (idx == first_heading_idx)   # đường cũ (không AI): heading đầu tiên = tiêu đề chính
            est_size = item.get("font_size")
            if is_title:
                m = DOC_TYPE_SPLIT_PATTERN.match(text)
                if m:
                    doc_type, _, rest = m.groups()
                    _add_para(doc, doc_type.strip(), align=WD_ALIGN_PARAGRAPH.CENTER, size=est_size or 14,
                              bold=True, space_after=2)
                    trich_yeu = ("Về việc " + rest.strip()) if rest.strip() else doc_type.strip()
                    _add_para(doc, trich_yeu, align=WD_ALIGN_PARAGRAPH.CENTER, size=est_size or 13, bold=True,
                              space_after=10)
                else:
                    _add_para(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, size=est_size or 14, bold=True,
                              space_after=10)
            else:
                # tiêu đề mục trong thân bài (1., 2., II., IV. ...) -> căn trái
                level = item.get("level", 1)
                default_size = 14 if level <= 1 else 13
                _add_para(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT, size=est_size or default_size, bold=True,
                          space_before=8, space_after=6)

        elif kind == "paragraph":
            text = item["text"]
            align_hint = item.get("align")
            italic = bool(item.get("italic", False))
            size = item.get("font_size") or 13
            if align_hint == "center" or (align_hint is None and KINHGUI_PATTERN.match(text)):
                _add_para(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, size=size, italic=italic,
                          space_after=10)
            elif align_hint == "left":
                _add_para(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT, size=size, italic=italic,
                          space_after=8)
            else:
                _add_para(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=size, italic=italic,
                          first_line_indent=1.0, space_after=8)

    doc.save(output_path)


def render_preview_into_widget(text_widget, render_items):
    """Hiển thị bản xem trước (không phải file .docx thật) vào 1 Tkinter Text
    widget, dùng tag để mô phỏng đậm/nghiêng/căn giữa - đủ để xem lại nội dung
    trước khi xuất file Word thật."""
    text_widget.config(state="normal")
    text_widget.delete("1.0", "end")

    text_widget.tag_configure("bold", font=("Times New Roman", 12, "bold"))
    text_widget.tag_configure("italic", font=("Times New Roman", 11, "italic"))
    text_widget.tag_configure("center", justify="center")
    text_widget.tag_configure("title", font=("Times New Roman", 13, "bold"), justify="center")
    text_widget.tag_configure("section", font=("Times New Roman", 12, "bold"))
    text_widget.tag_configure("normal", font=("Times New Roman", 11))
    text_widget.tag_configure("table_hdr", font=("Courier New", 10, "bold"))
    text_widget.tag_configure("table_cell", font=("Courier New", 10))

    first_heading_idx = next((i for i, it in enumerate(render_items) if it["kind"] == "heading"), None)

    for idx, item in enumerate(render_items):
        kind = item["kind"]

        if kind == "letterhead":
            org = " - ".join(item.get("org_lines") or [])
            line = org
            if item.get("so_hieu"):
                line += ("\n" if line else "") + item["so_hieu"]
            right = item.get("quoc_hieu", "")
            if item.get("tieu_ngu"):
                right += "\n" + item["tieu_ngu"]
            if item.get("ngay_thang"):
                right += "\n" + item["ngay_thang"]
            text_widget.insert("end", (line or "(không rõ cơ quan)") + "\n", ("bold", "center"))
            text_widget.insert("end", right + "\n\n", ("italic", "center"))

        elif kind == "signature_block":
            text_widget.insert("end", "Nơi nhận:\n", ("italic",))
            rest = NOINHAN_PATTERN.sub("", item.get("noi_nhan_text", "")).strip()
            items = item.get("noi_nhan_items")
            if items is None:
                items = [p.strip(" ;.") for p in re.split(r"\s*-\s+", rest) if p.strip(" ;.")]
            for line in items:
                text_widget.insert("end", f"  - {line}\n", ("italic",))
            if item.get("signature_text"):
                text_widget.insert("end", "\n" + item["signature_text"] + "\n", ("bold", "center"))
            text_widget.insert("end", "\n")

        elif kind == "table":
            rows = item.get("rows") or []
            if rows:
                widths = [max(len(str(r[c])) if c < len(r) else 0 for r in rows)
                          for c in range(max(len(r) for r in rows))]
                for r_idx, row in enumerate(rows):
                    cells = [str(row[c]) if c < len(row) else "" for c in range(len(widths))]
                    line = " | ".join(c.ljust(w) for c, w in zip(cells, widths))
                    tag = "table_hdr" if r_idx == 0 else "table_cell"
                    text_widget.insert("end", line + "\n", (tag,))
                text_widget.insert("end", "\n")

        elif kind == "heading":
            is_title = item.get("is_title")
            if is_title is None:
                is_title = (idx == first_heading_idx)
            tag = "title" if is_title else "section"
            text_widget.insert("end", item["text"] + "\n\n", (tag,))

        elif kind == "paragraph":
            align = item.get("align")
            tags = ["normal"]
            if item.get("italic"):
                tags = ["italic"]
            if align == "center" or (align is None and KINHGUI_PATTERN.match(item["text"])):
                tags.append("center")
            text_widget.insert("end", item["text"] + "\n\n", tuple(tags))

    text_widget.config(state="disabled")


AI_MODEL = "claude-sonnet-4-6"


# ----------------------------------------------------------------------------
# CÁC ENGINE NẶNG: chỉ load khi thực sự bắt đầu chuyển đổi (lazy), tránh app
# mở chậm / treo khi chưa cần dùng.
# ----------------------------------------------------------------------------
_vietocr_predictor = None
_ppstructure_engine = None


def get_vietocr():
    global _vietocr_predictor
    if _vietocr_predictor is None:
        from vietocr.tool.predictor import Predictor
        from vietocr.tool.config import Cfg
        cfg = Cfg.load_config_from_name("vgg_transformer")
        cfg["device"] = "cpu"   # đổi thành "cuda:0" nếu máy có GPU NVIDIA hỗ trợ CUDA
        # Nếu có sẵn file weights "vgg_transformer.pth" ĐÃ TẢI SẴN đóng gói kèm
        # app (trong thư mục con "models") -> dùng luôn file đó, máy nào chạy
        # cũng KHÔNG cần tải lại từ Internet ở lần đầu chạy. Nếu không thấy thì
        # giữ hành vi cũ (để VietOCR tự tải về như trước).
        for base in _bundled_dir_candidates("models"):
            bundled_weights = os.path.join(base, "vgg_transformer.pth")
            if os.path.isfile(bundled_weights):
                cfg["weights"] = bundled_weights
                break
        _vietocr_predictor = Predictor(cfg)
    return _vietocr_predictor


def get_engine():
    global _ppstructure_engine
    if _ppstructure_engine is None:
        from paddleocr import PPStructureV3
        # KHÔNG đặt enable_mkldnn=False: cách đó né được lỗi PIR/oneDNN của
        # PaddlePaddle 3.3.x nhưng lại gây lỗi rò bộ nhớ rất nặng khi predict()
        # (có báo cáo ăn tới 43GB RAM, làm máy treo rồi bị buộc tắt). Thay vào
        # đó: hạ paddlepaddle về bản 3.2.2 (pip install paddlepaddle==3.2.2) -
        # bản này không có cả hai lỗi trên, giữ mkldnn mặc định (bật) để chạy
        # nhanh và ổn định.
        _ppstructure_engine = PPStructureV3(
            lang="vi",
            use_doc_orientation_classify=False,
            use_doc_unwarping=False,
        )
    return _ppstructure_engine


# ----------------------------------------------------------------------------
# HÀM TIỆN ÍCH: cắt ảnh, nhận diện 1 dòng, nhận diện cả khối theo toạ độ
# ----------------------------------------------------------------------------
def suppress_light_watermark(img_bgr, brightness_threshold=190):
    """Watermark/dấu mờ (kiểu chữ 'VATM' to mờ hoặc chữ chìm lặp chéo trang) thường
    in màu xám rất nhạt, sáng hơn nhiều so với chữ thật (đen/đậm). Xoá các pixel quá
    sáng thành trắng tinh trước khi đưa qua OCR, giảm khả năng model đọc nhầm
    watermark thành nội dung. Chữ thật (đen, đậm) không bị ảnh hưởng vì luôn tối
    hơn ngưỡng này nhiều."""
    gray = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2GRAY)
    mask = gray > brightness_threshold
    img_bgr[mask] = 255
    return img_bgr


_GARBAGE_WORD_PATTERN = re.compile(r"[A-Za-z]{10,}")
_VIET_DIACRITIC_PATTERN = re.compile(
    "[ăâđêôơưĂÂĐÊÔƠƯáàảãạấầẩẫậắằẳẵặéèẻẽẹếềểễệíìỉĩịóòỏõọốồổỗộớờởỡợúùủũụứừửữựýỳỷỹỵ"
    "ÁÀẢÃẠẤẦẨẪẬẮẰẲẴẶÉÈẺẼẸẾỀỂỄỆÍÌỈĨỊÓÒỎÕỌỐỒỔỖỘỚỜỞỠỢÚÙỦŨỤỨỪỬỮỰÝỲỶỸỴ]"
)


def looks_like_ocr_garbage(text):
    """Watermark mờ/dấu mộc đôi khi vẫn bị model đọc nhầm thành 'chữ', thường ra
    những từ tiếng Anh dài bất thường (ví dụ 'CONTRACTIONALIZED') hoặc chuỗi số
    lạ kiểu mã đăng ký trên dấu mộc - không phải nội dung văn bản thật. Không lọc
    chữ có dấu tiếng Việt (an toàn cho nội dung thật), chỉ nghi ngờ khi HOÀN TOÀN
    không có dấu tiếng Việt và có từ tiếng Anh rất dài kiểu này."""
    if _VIET_DIACRITIC_PATTERN.search(text):
        return False
    return bool(_GARBAGE_WORD_PATTERN.search(text))


def crop_region(img_bgr, box, pad=2):
    h, w = img_bgr.shape[:2]
    x0, y0, x1, y1 = box
    x0 = max(0, int(x0) - pad)
    y0 = max(0, int(y0) - pad)
    x1 = min(w, int(x1) + pad)
    y1 = min(h, int(y1) + pad)
    if x1 <= x0 or y1 <= y0:
        return None
    crop_bgr = img_bgr[y0:y1, x0:x1]
    return Image.fromarray(cv2.cvtColor(crop_bgr, cv2.COLOR_BGR2RGB))


def recognize_line(img_bgr, box, vietocr_predictor):
    pil_img = crop_region(img_bgr, box)
    if pil_img is None:
        return ""
    try:
        return vietocr_predictor.predict(pil_img).strip()
    except Exception as ex:
        print("VietOCR error:", ex)
        return ""


def box_center(box):
    x0, y0, x1, y1 = box
    return (x0 + x1) / 2, (y0 + y1) / 2


def center_inside(box, container_box, pad=10):
    # pad: nới biên một chút, vì bbox của khối layout đôi khi tính hẹp hơn 1-2 dòng
    # so với thực tế, khiến dòng chữ ở sát mép bị rơi ra ngoài và bị bỏ sót.
    cx, cy = box_center(box)
    bx0, by0, bx1, by1 = container_box
    return (bx0 - pad) <= cx <= (bx1 + pad) and (by0 - pad) <= cy <= (by1 + pad)


DPI = 300   # PHẢI khớp với dpi= trong convert_from_path() ở process_pdf, dùng để quy đổi px -> pt
COMMON_FONT_SIZES_PT = [8, 9, 10, 10.5, 11, 12, 13, 14, 15, 16, 18, 20, 22, 24, 26, 28, 32, 36]


def estimate_font_size_pt(line_heights_px, default=13):
    """Ước lượng cỡ chữ (pt) từ chiều cao dòng chữ trong ảnh đã OCR (đơn vị pixel,
    ở DPI render cố định). Hệ số 0.72 là hiệu chỉnh thực nghiệm: bbox 1 dòng chữ
    (tính cả phần đuôi/mũ ký tự như 'g', 'ệ'...) thường cao hơn cỡ chữ danh nghĩa
    (point size) một khoảng nhất định. Sau khi quy đổi, chọn cỡ chữ THÔNG DỤNG
    gần nhất (làm tròn về các cỡ hay dùng trong Word) để tránh ra số lẻ kiểu
    13.37pt - vừa không chuyên nghiệp, vừa làm các dòng cùng cỡ chữ thật bị lệch
    nhau do nhiễu nhỏ khi đo bbox."""
    if not line_heights_px:
        return default
    heights = sorted(line_heights_px)
    median_h = heights[len(heights) // 2]
    raw_pt = median_h * 72.0 / DPI * 0.72
    return min(COMMON_FONT_SIZES_PT, key=lambda s: abs(s - raw_pt))


def recognize_block_text(img_bgr, block_box, line_boxes, vietocr_predictor, used_mask=None, line_indices=None):
    if line_indices is None:
        line_indices = range(len(line_boxes))
    matched = [i for i in line_indices if center_inside(line_boxes[i], block_box)]
    matched.sort(key=lambda i: (line_boxes[i][1], line_boxes[i][0]))   # trên -> dưới, trái -> phải
    if used_mask is not None:
        for i in matched:
            used_mask[i] = True
    texts = [recognize_line(img_bgr, line_boxes[i], vietocr_predictor) for i in matched]
    text = " ".join(t for t in texts if t)
    heights = [line_boxes[i][3] - line_boxes[i][1] for i in matched]
    font_size = estimate_font_size_pt(heights) if text else None
    return text, font_size


def group_orphan_lines(img_bgr, orphan_indices, line_boxes, vietocr_predictor, gap_factor=1.6):
    """Các dòng không khớp bbox của bất kỳ khối nào (do bbox tính hẹp, hoặc khối
    không xác định được loại) -> vẫn nhận diện và gom lại thành đoạn văn theo vị
    trí, để không bị mất nội dung, thay vì âm thầm bỏ qua."""
    if not orphan_indices:
        return []
    ordered = sorted(orphan_indices, key=lambda i: (line_boxes[i][1], line_boxes[i][0]))
    heights = [line_boxes[i][3] - line_boxes[i][1] for i in ordered]
    avg_h = sum(heights) / len(heights) if heights else 20

    paragraphs = []
    current = [ordered[0]]
    for prev_i, cur_i in zip(ordered, ordered[1:]):
        gap = line_boxes[cur_i][1] - line_boxes[prev_i][3]
        if gap > avg_h * gap_factor:
            paragraphs.append(current)
            current = []
        current.append(cur_i)
    if current:
        paragraphs.append(current)

    result = []
    for group in paragraphs:
        texts = [recognize_line(img_bgr, line_boxes[i], vietocr_predictor) for i in group]
        text = " ".join(t for t in texts if t)
        if text:
            heights_g = [line_boxes[i][3] - line_boxes[i][1] for i in group]
            result.append((text, estimate_font_size_pt(heights_g)))
    return result


def _cluster_boxes_by_row(cell_boxes, n_rows):
    """Gom các toạ độ ô theo TỪNG DÒNG. Thay vì dùng 1 ngưỡng khoảng cách cố định
    (dễ sai khi các dòng cao thấp rất khác nhau trong cùng 1 bảng), cắt tại đúng
    (n_rows - 1) khoảng hở LỚN NHẤT giữa các ô đã sắp theo cạnh trên (top) - cách
    này LUÔN cho ra đúng n_rows nhóm (miễn có ít nhất n_rows ô), không phụ thuộc
    việc đoán ngưỡng có khớp hay không."""
    if not cell_boxes or n_rows <= 0 or len(cell_boxes) < n_rows:
        return None
    order = sorted(range(len(cell_boxes)), key=lambda i: cell_boxes[i][1])
    if n_rows == 1:
        return [order]
    gaps = [(cell_boxes[order[k + 1]][1] - cell_boxes[order[k]][1], k) for k in range(len(order) - 1)]
    gaps.sort(key=lambda t: -t[0])
    cut_after = sorted(k for _, k in gaps[: n_rows - 1])
    groups, start = [], 0
    for k in cut_after:
        groups.append(order[start:k + 1])
        start = k + 1
    groups.append(order[start:])
    return groups


def _compute_column_bands(cell_boxes, block_box, n_cols, iterations=12):
    """Xác định vị trí X đại diện cho từng CỘT bằng K-means trên toạ độ X của tất
    cả ô phát hiện được, khởi tạo tâm cụm ở vị trí CHIA ĐỀU theo chiều rộng bảng
    (block_box). Nếu 1 cột không có ô nào (ví dụ cột STT quá hẹp, model không
    nhận diện được bbox ở hầu hết/mọi dòng), tâm cụm của cột đó không bị kéo đi
    đâu cả, vẫn giữ đúng vị trí chia đều ban đầu - vẫn là một vị trí hợp lý để so
    khớp, không làm hỏng việc xác định các cột khác."""
    if not cell_boxes or block_box is None or n_cols <= 0:
        return None
    bx0, _, bx1, _ = block_box
    centers = [bx0 + (c + 0.5) * (bx1 - bx0) / n_cols for c in range(n_cols)]
    xs = [(b[0] + b[2]) / 2 for b in cell_boxes]

    for _ in range(iterations):
        buckets = [[] for _ in range(n_cols)]
        for x in xs:
            nearest = min(range(n_cols), key=lambda c: abs(x - centers[c]))
            buckets[nearest].append(x)
        new_centers = [(sum(b) / len(b)) if b else centers[c] for c, b in enumerate(buckets)]
        if new_centers == centers:
            break
        centers = new_centers
    return centers


def _assign_row_boxes_to_columns(row_box_indices, cell_boxes, column_bands):
    """Gán mỗi ô trong 1 dòng vào đúng CỘT theo toạ độ X. Dùng quy hoạch động
    (DP) để đảm bảo tính chất BẮT BUỘC: ô càng bên trái luôn được gán vào cột có
    số thứ tự nhỏ hơn hoặc bằng ô bên phải nó - KHÔNG BAO GIỜ đảo thứ tự cột dù
    vị trí cột (column_bands) ước lượng chưa hoàn toàn chuẩn. Đây là lỗi đã gặp
    trước đây: cách chọn "cột gần nhất" cho từng ô độc lập (không ràng buộc thứ
    tự) có thể vô tình gán 1 ô vào cột nằm GIỮA bảng dù ô đó thực ra ở đầu/cuối
    dòng, làm cả dòng bị đảo lộn. Trả về dict {col_idx: box_idx}."""
    if not row_box_indices:
        return {}
    n_cols = len(column_bands)
    boxes_sorted = sorted(row_box_indices, key=lambda i: (cell_boxes[i][0] + cell_boxes[i][2]) / 2)
    if len(boxes_sorted) > n_cols:
        # Nhieu o hon so cot (hiem, co the do nhieu/tach sai) - chi giu dung
        # n_cols o gan tam cac cot nhat de DP van chay duoc.
        boxes_sorted = boxes_sorted[:n_cols]
    xs = [(cell_boxes[i][0] + cell_boxes[i][2]) / 2 for i in boxes_sorted]
    k = len(xs)

    INF = float("inf")
    # dp[i][j] = chi phí nhỏ nhất khi đã gán xong i ô đầu (ô thứ i, chỉ số i-1,
    # gán vào đúng cột j); choice[i][j] = cột đã dùng cho ô (i-1) để đạt chi phí đó.
    dp = [[INF] * n_cols for _ in range(k + 1)]
    choice = [[-1] * n_cols for _ in range(k + 1)]
    for j in range(n_cols):
        dp[1][j] = abs(xs[0] - column_bands[j])
    for i in range(2, k + 1):
        best_so_far, best_jp = INF, -1
        for j in range(n_cols):
            # best_so_far = min(dp[i-1][0..j-1]) - tính dồn để cả hàm chạy O(k*n_cols)
            if j > 0 and dp[i - 1][j - 1] < best_so_far:
                best_so_far, best_jp = dp[i - 1][j - 1], j - 1
            if best_jp != -1:
                dp[i][j] = best_so_far + abs(xs[i - 1] - column_bands[j])
                choice[i][j] = best_jp

    best_j = min(range(n_cols), key=lambda j: dp[k][j])
    if dp[k][best_j] == INF:
        return {}

    assign_cols = [0] * k
    j = best_j
    for i in range(k, 0, -1):
        assign_cols[i - 1] = j
        j = choice[i][j]

    return {assign_cols[idx]: boxes_sorted[idx] for idx in range(k)}


def _representative_font_size(sizes, default=12):
    sizes = [s for s in sizes if s]
    if not sizes:
        return default
    counts = {}
    for s in sizes:
        counts[s] = counts.get(s, 0) + 1
    max_count = max(counts.values())
    candidates = sorted(s for s, c in counts.items() if c == max_count)
    return candidates[len(candidates) // 2]


def build_table_rows(img_bgr, table_res_list, table_index, line_boxes, vietocr_predictor, log,
                      used_mask=None, block_box=None):
    if table_index >= len(table_res_list):
        return None
    table_res = table_res_list[table_index]
    html = table_res.get("pred_html", "")
    cell_boxes = table_res.get("cell_box_list", [])

    soup = BeautifulSoup(html, "html.parser")
    rows = soup.find_all("tr")
    if not rows:
        return None

    all_cells = [c for r in rows for c in r.find_all(["td", "th"])]
    ncols_guess = max(len(r.find_all(["td", "th"])) for r in rows)
    collected_sizes = []

    # CÁCH CHÍNH: gom toạ độ theo từng dòng (trục Y), rồi trong mỗi dòng khớp
    # từng ô vào đúng CỘT theo toạ độ X (không theo thứ tự còn lại). An toàn hơn
    # hẳn so với đánh số tuần tự toàn bảng - nếu 1 dòng/1 ô bị thiếu toạ độ thì
    # chỉ đúng ô đó dùng chữ HTML gốc, không kéo lệch dây chuyền các ô/dòng khác
    # (lỗi đã gặp: số liệu/ô bị "trượt" dồn sang cột/dòng khác).
    row_groups = _cluster_boxes_by_row(cell_boxes, len(rows))
    column_bands = _compute_column_bands(cell_boxes, block_box, ncols_guess)

    if row_groups is not None and column_bands is not None:
        log(f"  [Bảng] Khớp toạ độ theo đúng dòng + đúng cột ({len(rows)} dòng x {ncols_guess} cột) "
            f"- tránh lệch dây chuyền.")
        result_rows = []
        col_widths_px = [[] for _ in range(ncols_guess)]
        for r_idx, row in enumerate(rows):
            cols = row.find_all(["td", "th"])
            assignment = _assign_row_boxes_to_columns(row_groups[r_idx], cell_boxes, column_bands)
            row_out = []
            for c_idx, col in enumerate(cols):
                box_i = assignment.get(c_idx)
                if box_i is not None:
                    if c_idx < ncols_guess:
                        col_widths_px[c_idx].append(cell_boxes[box_i][2] - cell_boxes[box_i][0])
                    txt, sz = recognize_block_text(img_bgr, cell_boxes[box_i], line_boxes, vietocr_predictor,
                                                     used_mask=used_mask)
                    if sz:
                        collected_sizes.append(sz)
                    if not txt:
                        txt = col.get_text(strip=True)
                else:
                    txt = col.get_text(strip=True)
                row_out.append(txt)
            result_rows.append(row_out)

        # Tỉ lệ độ rộng cột giống bản gốc: dùng độ rộng ô THẬT đo được cho cột đã
        # có dữ liệu; cột nào không có ô nào được phát hiện (thường vì quá hẹp,
        # như cột STT) thì đoán là HẸP HƠN cột hẹp nhất đã đo được (lý do phổ biến
        # khiến model không nhận ra bbox chính là cột đó quá hẹp), thay vì đoán
        # bằng độ rộng trung bình (sẽ làm cột đó bị xuất ra quá rộng so với thật).
        avg_widths = [sum(w) / len(w) for w in col_widths_px if w]
        unmeasured_default = (min(avg_widths) * 0.5) if avg_widths else None
        final_widths = []
        for w in col_widths_px:
            if w:
                final_widths.append(sum(w) / len(w))
            elif unmeasured_default is not None:
                final_widths.append(unmeasured_default)
            else:
                final_widths.append(None)
        col_width_ratios = None
        if all(w is not None for w in final_widths) and sum(final_widths) > 0:
            total_w = sum(final_widths)
            col_width_ratios = [w / total_w for w in final_widths]

        return result_rows, _representative_font_size(collected_sizes), col_width_ratios

    if row_groups is not None:
        # Gom được theo dòng nhưng không xác định được cột -> ghép tuần tự trong
        # từng dòng (vẫn an toàn hơn đánh số toàn bảng, vì không lan sang dòng khác).
        log(f"  [Bảng] Ghép toạ độ theo từng dòng ({len(rows)} dòng), không xác định được vị trí cột "
            f"-> ghép tuần tự trong từng dòng.")
        result_rows = []
        for r_idx, row in enumerate(rows):
            cols = row.find_all(["td", "th"])
            group = sorted(row_groups[r_idx], key=lambda i: cell_boxes[i][0])
            row_out = []
            for c_idx, col in enumerate(cols):
                if c_idx < len(group):
                    box = cell_boxes[group[c_idx]]
                    txt, sz = recognize_block_text(img_bgr, box, line_boxes, vietocr_predictor, used_mask=used_mask)
                    if sz:
                        collected_sizes.append(sz)
                    if not txt:
                        txt = col.get_text(strip=True)
                else:
                    txt = col.get_text(strip=True)
                row_out.append(txt)
            result_rows.append(row_out)
        return result_rows, _representative_font_size(collected_sizes), None

    # --- KHÔNG gom được theo dòng (số nhóm khác số dòng HTML) -> phương án dự phòng cũ ---
    log(f"  [Bảng] Không gom được toạ độ theo từng dòng (số dòng HTML={len(rows)}) "
        f"-> dùng phương án dự phòng (đánh số tuần tự / chia lưới đều).")

    # Nếu thiếu toạ độ ô quá nhiều (ví dụ model không trả về cell_box_list) ->
    # tự chia đều vùng bảng (bbox của khối layout) theo lưới số dòng x số cột,
    # vẫn còn hơn là bỏ hẳn VietOCR cho cả bảng.
    if len(cell_boxes) < len(all_cells) * 0.5 and block_box is not None:
        bx0, by0, bx1, by1 = block_box
        row_h = (by1 - by0) / max(len(rows), 1)
        col_w = (bx1 - bx0) / max(ncols_guess, 1)
        grid_boxes = []
        for r_idx in range(len(rows)):
            for c_idx in range(ncols_guess):
                grid_boxes.append([
                    bx0 + c_idx * col_w, by0 + r_idx * row_h,
                    bx0 + (c_idx + 1) * col_w, by0 + (r_idx + 1) * row_h,
                ])
        if len(grid_boxes) > len(cell_boxes):
            log(f"  [Bảng] Thiếu toạ độ ô (chỉ có {len(cell_boxes)}/{len(all_cells)}) -> "
                f"tự chia đều vùng bảng theo lưới {len(rows)}x{ncols_guess} để vẫn dùng được VietOCR.")
            cell_boxes = grid_boxes
    elif len(all_cells) != len(cell_boxes):
        log(f"  [Bảng] Số ô ({len(all_cells)}) không khớp số toạ độ ô ({len(cell_boxes)}) "
            f"-> {min(len(all_cells), len(cell_boxes))} ô đầu vẫn dùng VietOCR, "
            f"phần dư (nếu có) mới giữ chữ gốc.")

    result_rows = []
    cell_i = 0
    for row in rows:
        cols = row.find_all(["td", "th"])
        row_out = []
        for col in cols:
            # Chỉ dùng chữ gốc (model "latin", thiếu dấu tiếng Việt) cho đúng những ô
            # KHÔNG có toạ độ tương ứng -> không vì 1-2 ô lệch mà bỏ VietOCR cho cả bảng.
            if cell_i < len(cell_boxes):
                txt, sz = recognize_block_text(img_bgr, cell_boxes[cell_i], line_boxes, vietocr_predictor,
                                                used_mask=used_mask)
                if sz:
                    collected_sizes.append(sz)
                if not txt:
                    txt = col.get_text(strip=True)
            else:
                txt = col.get_text(strip=True)
            row_out.append(txt)
            cell_i += 1
        result_rows.append(row_out)
    return result_rows, _representative_font_size(collected_sizes), None


# ----------------------------------------------------------------------------
# CHỌN TRANG: "1-3,5,8-10" -> [1,2,3,5,8,9,10]
# ----------------------------------------------------------------------------
def parse_page_selection(text, total_pages):
    text = (text or "").strip()
    if not text:
        raise ValueError("Chưa nhập trang nào.")
    pages = set()
    for part in text.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            a, b = part.split("-", 1)
            a, b = int(a.strip()), int(b.strip())
            pages.update(range(min(a, b), max(a, b) + 1))
        else:
            pages.add(int(part))
    result = sorted(p for p in pages if 1 <= p <= total_pages)
    if not result:
        raise ValueError("Không có trang hợp lệ trong khoảng đã chọn.")
    return result


# ----------------------------------------------------------------------------
# (TUỲ CHỌN) GỌI CLAUDE API ĐỂ DỰNG LẠI TOÀN BỘ CẤU TRÚC TRÌNH BÀY
# Thay cho việc chỉ sửa lỗi OCR + cấp tiêu đề: khi bật AI, để AI tự nhận diện
# luôn cả khối quốc hiệu/cơ quan, tiêu đề chính, "Kính gửi", tiêu đề mục, khối
# "Nơi nhận"/ký tên - giống người đọc hiểu văn bản, nên xử lý được cả các
# trường hợp bất thường (bbox lệch, thứ tự bị xáo) mà cách dò bằng mẫu cố định
# (restructure_document/restructure_signature_blocks) không xử lý được.
# Bảng được gửi dưới dạng placeholder (không gửi nội dung) để AI giữ đúng vị
# trí, dữ liệu bảng thật luôn lấy lại từ render_items gốc, AI không có cơ hội
# làm sai lệch nội dung bảng.
# ----------------------------------------------------------------------------
def apply_ai_restructure(api_key, render_items, log):
    import anthropic

    n = len(render_items)
    payload = []
    for i, item in enumerate(render_items):
        if item["kind"] == "table":
            payload.append({"i": i, "text": f"[BẢNG #{i} - chỉ là placeholder để giữ vị trí, không sửa]"})
        else:
            payload.append({"i": i, "text": item["text"]})

    prompt = (
        "Bạn nhận một danh sách các khối nội dung (đoạn văn/tiêu đề/bảng) được trích "
        "xuất bằng OCR từ một văn bản hành chính/pháp luật tiếng Việt, theo thứ tự đọc "
        "(có thể có vài khối bị lệch thứ tự do lỗi nhận diện vị trí - bạn cần tự nhận "
        "ra và sắp lại đúng). Mỗi khối có 'i' (số thứ tự gốc) và 'text'.\n\n"
        "Hãy dựng lại cấu trúc trình bày CHUẨN của văn bản hành chính Việt Nam, gộp các "
        "khối liên quan với nhau, trả về một JSON array các 'block', mỗi block là MỘT "
        "trong các dạng sau:\n\n"
        '1. {"type":"letterhead","consumes":[...],"org_lines":["tên cơ quan, có thể '
        'nhiều dòng"],"so_hieu":"Số: .../...","quoc_hieu":"CỘNG HÒA XÃ HỘI CHỦ NGHĨA '
        'VIỆT NAM","tieu_ngu":"Độc lập - Tự do - Hạnh phúc","ngay_thang":"Hà Nội, ngày '
        '... tháng ... năm ..."}  - khối quốc hiệu/cơ quan ban hành, thường 1 lần đầu '
        "mỗi văn bản con.\n"
        '2. {"type":"title","consumes":[i],"doc_type":"TỜ TRÌNH","trich_yeu":"Về việc '
        '..."} - tên loại văn bản + trích yếu, CHỈ 1 lần đầu mỗi văn bản con.\n'
        '3. {"type":"kinh_gui","consumes":[i],"text":"Kính gửi: ..."}\n'
        '4. {"type":"section_heading","consumes":[i],"text":"...","level":1 hoặc 2} - '
        "tiêu đề mục trong thân bài, đánh số kiểu \"1.\", \"2.\", \"II.\", \"Điều 1.\".\n"
        '5. {"type":"signature_block","consumes":[...],"noi_nhan_items":["Như trên", '
        '"HĐTV (để b/c)", ...],"signature_text":"KT. ...\\nPHÓ TRƯỞNG BAN\\n\\nTên người '
        'ký"} - khối "Nơi nhận:" gộp với chức danh/tên người ký gần đó.\n'
        '6. {"type":"table","consumes":[i]} - giữ placeholder bảng đúng vị trí, không '
        "thêm field khác.\n"
        '7. {"type":"paragraph","consumes":[i],"text":"...","align":"justify" hoặc '
        '"center","italic":true/false} - đoạn văn thường; dùng "center"+italic cho các '
        'chú thích ngắn kiểu "(Chi tiết tại ... kèm theo)".\n\n'
        "QUY TẮC BẮT BUỘC:\n"
        "- Sửa lỗi OCR/chính tả rõ ràng trong 'text', KHÔNG đoán bừa nếu không chắc, "
        "không thêm nội dung không có trong văn bản gốc, không tự dịch/diễn giải lại.\n"
        "- 'consumes' phải liệt kê đúng, đủ các 'i' gốc được gộp vào block đó; mỗi 'i' "
        "chỉ được xuất hiện trong ĐÚNG 1 block, không bỏ sót, không lặp.\n"
        "- Với block 'table': không sửa, không thêm field nội dung.\n"
        "- Giữ đúng trình tự logic của văn bản gốc.\n"
        "- Nếu không chắc 1 khối thuộc loại nào, dùng 'paragraph' với align 'justify' "
        "(an toàn nhất, không tự bịa loại).\n\n"
        "Trả về DUY NHẤT JSON array, không kèm giải thích, không markdown.\n\n"
        f"Dữ liệu đầu vào:\n{json.dumps(payload, ensure_ascii=False)}"
    )

    client = anthropic.Anthropic(api_key=api_key)
    resp = client.messages.create(
        model=AI_MODEL,
        max_tokens=8000,
        messages=[{"role": "user", "content": prompt}],
    )
    raw = "".join(b.text for b in resp.content if getattr(b, "type", "") == "text").strip()
    raw = re.sub(r"^```[a-zA-Z]*\n|```$", "", raw, flags=re.MULTILINE).strip()
    blocks = json.loads(raw)

    covered = set()
    new_items = []
    for b in blocks:
        consumes = [i for i in (b.get("consumes") or []) if 0 <= i < n and i not in covered]
        if not consumes:
            continue
        for i in consumes:
            covered.add(i)
        page = render_items[consumes[0]].get("page")
        font_size = render_items[consumes[0]].get("font_size")
        btype = b.get("type")

        if btype == "table":
            orig = render_items[consumes[0]]
            if orig["kind"] == "table":
                new_items.append({"kind": "table", "rows": orig["rows"], "page": page,
                                   "font_size": orig.get("font_size"),
                                   "col_width_ratios": orig.get("col_width_ratios")})

        elif btype == "letterhead":
            new_items.append({
                "kind": "letterhead", "page": page,
                "org_lines": b.get("org_lines") or [],
                "so_hieu": b.get("so_hieu", ""),
                "quoc_hieu": b.get("quoc_hieu") or QUOCHIEU_TEXT,
                "tieu_ngu": b.get("tieu_ngu", ""),
                "ngay_thang": b.get("ngay_thang", ""),
            })

        elif btype == "title":
            doc_type = (b.get("doc_type") or "").strip()
            trich_yeu = (b.get("trich_yeu") or "").strip()
            text = (doc_type + (" " + trich_yeu if trich_yeu else "")).strip()
            new_items.append({"kind": "heading", "text": text or "(không có tiêu đề)",
                               "is_title": True, "level": 1, "page": page, "font_size": font_size})

        elif btype == "kinh_gui":
            new_items.append({"kind": "paragraph", "text": b.get("text", ""),
                               "align": "center", "page": page, "font_size": font_size})

        elif btype == "section_heading":
            new_items.append({"kind": "heading", "text": b.get("text", ""), "is_title": False,
                               "level": int(b.get("level", 1)) or 1, "page": page, "font_size": font_size})

        elif btype == "signature_block":
            new_items.append({
                "kind": "signature_block", "page": page,
                "noi_nhan_items": b.get("noi_nhan_items") or [],
                "signature_text": b.get("signature_text", ""),
            })

        else:   # "paragraph" hoặc loại lạ AI tự đặt ra -> render an toàn như đoạn văn
            new_items.append({"kind": "paragraph", "text": b.get("text", ""),
                               "align": b.get("align", "justify"),
                               "italic": bool(b.get("italic", False)), "page": page, "font_size": font_size})

    missing = [i for i in range(n) if i not in covered]
    if missing:
        log(f"  AI bỏ sót {len(missing)}/{n} khối -> tự thêm lại ở cuối để không mất nội dung.")
        for i in missing:
            orig = render_items[i]
            if orig["kind"] == "table":
                new_items.append({"kind": "table", "rows": orig["rows"], "page": orig.get("page"),
                                   "font_size": orig.get("font_size"),
                                   "col_width_ratios": orig.get("col_width_ratios")})
            else:
                new_items.append({"kind": "paragraph", "text": orig["text"], "page": orig.get("page"),
                                   "font_size": orig.get("font_size")})

    return new_items


# ----------------------------------------------------------------------------
# (TUỲ CHỌN, CẦN AI) SỬA LỖI CHÍNH TẢ/THIẾU DẤU TRONG TỪNG Ô BẢNG
# apply_ai_restructure() ở trên CHỦ ĐỘNG không cho AI thấy nội dung bảng (chỉ
# gửi placeholder) để AI không làm hỏng cấu trúc bảng -> nghĩa là chữ trong
# bảng chưa từng được AI sửa lỗi, dù đoạn văn thường thì có. Hàm này bù lại:
# chỉ sửa CHỮ trong từng ô, không đổi số dòng/số cột, không đụng vào "rows".
# ----------------------------------------------------------------------------
def apply_ai_table_spellcheck(api_key, render_items, log):
    import anthropic

    cells_payload = []
    for item_idx, item in enumerate(render_items):
        if item["kind"] != "table":
            continue
        for r_idx, row in enumerate(item["rows"]):
            for c_idx, cell_text in enumerate(row):
                if cell_text and cell_text.strip():
                    cells_payload.append({"item": item_idx, "r": r_idx, "c": c_idx, "text": cell_text})

    if not cells_payload:
        return render_items

    prompt = (
        "Bạn nhận một danh sách các đoạn chữ ngắn trong ô bảng, trích xuất bằng OCR từ "
        "văn bản hành chính tiếng Việt, có thể còn lỗi chính tả/thiếu dấu do OCR (ví dụ: "
        "'Quyt đnh b hy bó' đúng phải là 'Quyết định bị hủy bỏ'; 'S tay cht lng' đúng "
        "phải là 'Sổ tay chất lượng'; 'son v cp bn' đúng phải là 'soạn và cập nhật bản'). "
        "Hãy sửa lại đúng chính tả/dấu tiếng Việt cho MỖI ô, GIỮ NGUYÊN ý nghĩa và số "
        "liệu/mã văn bản, KHÔNG dịch, không diễn giải lại, không thêm nội dung không có. "
        "Nếu một ô đã đúng hoặc không chắc cách sửa, giữ nguyên y như input.\n\n"
        "Trả về DUY NHẤT một JSON array, ĐÚNG số lượng và thứ tự như đầu vào, mỗi phần tử "
        'dạng {"item":..,"r":..,"c":..,"text":"<chữ đã sửa>"}\n\n'
        f"Dữ liệu:\n{json.dumps(cells_payload, ensure_ascii=False)}"
    )

    client = anthropic.Anthropic(api_key=api_key)
    resp = client.messages.create(
        model=AI_MODEL,
        max_tokens=8000,
        messages=[{"role": "user", "content": prompt}],
    )
    raw = "".join(b.text for b in resp.content if getattr(b, "type", "") == "text").strip()
    raw = re.sub(r"^```[a-zA-Z]*\n|```$", "", raw, flags=re.MULTILINE).strip()
    fixed = json.loads(raw)

    # QUAN TRỌNG: khớp lại theo đúng khoá (item, r, c) AI trả về, KHÔNG dùng zip()
    # theo thứ tự vị trí - AI có thể trả JSON array không đúng thứ tự gửi đi (đặc
    # biệt với bảng nhiều ô), zip theo vị trí sẽ gán nhầm chữ sang ô khác, gây hiện
    # tượng số/ô bị đảo lộn giữa các dòng.
    fixed_by_key = {}
    for out in fixed:
        try:
            key = (int(out["item"]), int(out["r"]), int(out["c"]))
            fixed_by_key[key] = out.get("text", "")
        except (KeyError, ValueError, TypeError):
            continue

    n_applied = 0
    for src in cells_payload:
        key = (src["item"], src["r"], src["c"])
        if key in fixed_by_key:
            try:
                render_items[src["item"]]["rows"][src["r"]][src["c"]] = fixed_by_key[key]
                n_applied += 1
            except Exception:
                pass
        # không tìm thấy đúng khoá -> giữ nguyên chữ gốc của ô đó, không đoán bừa

    if n_applied < len(cells_payload):
        log(f"  AI sửa được {n_applied}/{len(cells_payload)} ô; các ô còn lại giữ chữ gốc "
            f"(không khớp được khoá item/r/c).")
    return render_items


# ----------------------------------------------------------------------------
# XỬ LÝ CHÍNH: chạy trong thread riêng để không treo giao diện
# ----------------------------------------------------------------------------
def process_pdf(pdf_path, page_list, output_path, use_ai, api_key, log_queue):
    log_file_path = os.path.splitext(output_path)[0] + "_log.txt"
    log_file = open(log_file_path, "w", encoding="utf-8")

    def log(msg):
        log_queue.put(("status", msg))
        log_file.write(msg + "\n")
        log_file.flush()   # ghi xuống đĩa ngay, để dù app có crash vẫn còn lại trong file

    try:
        log(f"(Log chi tiết cũng được lưu tại: {log_file_path})")

        log("Đang tải model VietOCR...")
        vietocr_predictor = get_vietocr()
        log("Đã tải xong VietOCR. Đang tải model PP-StructureV3...")
        engine = get_engine()
        log("Đã tải xong PP-StructureV3.")

        render_items = []     # danh sách item theo ĐÚNG thứ tự cuối cùng sẽ ghi vào Word
        sidebar_notes = []
        page_orientation = {}   # page_no -> True nếu trang đó là giấy ngang (khổ rộng hơn cao)

        total = len(page_list)
        for i, page_no in enumerate(page_list, start=1):
            log(f"Đang xử lý trang {page_no} ({i}/{total})...")

            try:
                pages = convert_from_path(
                    pdf_path, dpi=300, poppler_path=POPPLER_PATH,
                    first_page=page_no, last_page=page_no,
                )
            except Exception as e:
                if "poppler" in str(e).lower() or "pdfinfo" in str(e).lower():
                    raise RuntimeError(
                        "Máy này chưa cài Poppler (chương trình phụ trợ để đọc file PDF, không "
                        "phải lỗi của app). Cách sửa: tải Poppler tại "
                        "https://github.com/oschwartz10612/poppler-windows , giải nén, rồi đặt "
                        "thư mục 'poppler' (chứa thư mục con Library\\bin) CẠNH file app này, "
                        "hoặc cài vào C:\\poppler rồi thử lại."
                    ) from e
                raise
            page_w, page_h = pages[0].size
            page_orientation[page_no] = page_w > page_h   # rộng hơn cao -> giấy ngang
            img_path = f"_tmp_page_{page_no}.jpg"
            pages[0].save(img_path, "JPEG")
            img_bgr = cv2.imread(img_path)
            img_bgr = suppress_light_watermark(img_bgr)

            log(f"  (trang {page_no}) Đang chạy PP-StructureV3 (layout)...")
            results = engine.predict(img_bgr)
            for res in results:
                blocks = res["parsing_res_list"]
                ocr_res = res["overall_ocr_res"]
                line_boxes = np.array(ocr_res["rec_boxes"]).tolist() if len(ocr_res["rec_boxes"]) else []
                table_res_list = res.get("table_res_list", [])
                table_index = 0
                used_mask = [False] * len(line_boxes)   # đánh dấu dòng nào đã được 1 khối/ô bảng "nhận"
                log(f"  (trang {page_no}) Tìm được {len(blocks)} khối, {len(line_boxes)} dòng chữ. "
                    f"Đang chạy VietOCR cho từng dòng...")

                for block in blocks:
                    label = block.label
                    block_box = block.bbox

                    if label in TABLE_LABELS:
                        table_result = build_table_rows(img_bgr, table_res_list, table_index, line_boxes,
                                                          vietocr_predictor, log, used_mask=used_mask,
                                                          block_box=block_box)
                        table_index += 1
                        if table_result:
                            rows_data, table_font_size, col_width_ratios = table_result
                            render_items.append({"kind": "table", "rows": rows_data, "page": page_no,
                                                  "font_size": table_font_size,
                                                  "col_width_ratios": col_width_ratios})
                        continue

                    content, content_font_size = recognize_block_text(img_bgr, block_box, line_boxes,
                                                                        vietocr_predictor, used_mask=used_mask)
                    if not content:
                        continue

                    if label in DROP_LABELS:
                        sidebar_notes.append(f"[trang {page_no} - {label}] {content}")
                        continue

                    if looks_like_ocr_garbage(content):
                        # nghi là watermark/dấu mộc bị đọc nhầm thành chữ (không có dấu
                        # tiếng Việt + có từ tiếng Anh dài bất thường) -> không đưa vào
                        # thân văn bản, chỉ ghi lại để xem nếu cần.
                        log(f"  (trang {page_no}) Bỏ qua 1 khối nghi là watermark/dấu mộc: "
                            f"\"{content[:60]}\"")
                        sidebar_notes.append(f"[trang {page_no} - nghi watermark] {content}")
                        continue

                    # header/header_image (quốc hiệu, ngày tháng...) và mọi khối còn lại
                    # (kể cả label lạ chưa biết) -> vẫn đưa vào văn bản như đoạn văn thường,
                    # đúng vị trí PP-StructureV3 đã sắp.
                    kind = "heading" if label in TITLE_LABELS else "paragraph"
                    render_items.append({"kind": kind, "level": 1, "text": content, "page": page_no,
                                          "font_size": content_font_size})

                # Dòng chữ nào không thuộc khối nào (bbox của khối tính hẹp hơn thực tế,
                # hoặc khối không rơi vào nhóm nào ở trên) -> vẫn vớt lại, không bỏ mất.
                orphan_indices = [i for i, used in enumerate(used_mask) if not used]
                if orphan_indices:
                    log(f"  (trang {page_no}) {len(orphan_indices)} dòng không khớp khối nào -> "
                        f"đang vớt lại để không mất nội dung.")
                    for text, text_font_size in group_orphan_lines(img_bgr, orphan_indices, line_boxes,
                                                                     vietocr_predictor):
                        if looks_like_ocr_garbage(text):
                            log(f"  (trang {page_no}) Bỏ qua 1 đoạn mồ côi nghi là watermark/dấu mộc: "
                                f"\"{text[:60]}\"")
                            sidebar_notes.append(f"[trang {page_no} - nghi watermark] {text}")
                            continue
                        render_items.append({"kind": "paragraph", "level": 1, "text": text, "page": page_no,
                                              "font_size": text_font_size})

            log(f"  (trang {page_no}) Xong.")
            os.remove(img_path)

        if use_ai:
            log("Đang gọi AI (Claude) để dựng lại toàn bộ cấu trúc trình bày...")
            try:
                render_items = apply_ai_restructure(api_key, render_items, log)
            except Exception as ex:
                log(f"Lỗi gọi AI ({ex}) -> dùng cách dựng cấu trúc thông thường (không AI).")
                render_items = restructure_document(render_items)
                render_items = restructure_signature_blocks(render_items)

            log("Đang gọi AI sửa lỗi chính tả/thiếu dấu trong các ô bảng...")
            try:
                render_items = apply_ai_table_spellcheck(api_key, render_items, log)
            except Exception as ex:
                log(f"Lỗi sửa chính tả bảng bằng AI ({ex}) -> giữ chữ trong bảng như cũ.")
        else:
            log("Đang dựng khung văn bản (quốc hiệu/cơ quan, tiêu đề, Kính gửi, nơi nhận/ký tên)...")
            render_items = restructure_document(render_items)
            render_items = restructure_signature_blocks(render_items)

        log("Hoàn tất nhận diện & dựng cấu trúc - đang hiển thị bản xem trước...")
        log_queue.put(("preview_ready", render_items, page_orientation, sidebar_notes))
    finally:
        log_file.close()


def process_pdf_safe(pdf_path, page_list, output_path, use_ai, api_key, log_queue):
    try:
        process_pdf(pdf_path, page_list, output_path, use_ai, api_key, log_queue)
    except Exception:
        log_queue.put(("error", traceback.format_exc()))


# ----------------------------------------------------------------------------
# LƯU CẤU HÌNH (API key) GIỮA CÁC LẦN MỞ APP
# Lưu ý: file lưu dạng JSON thường (KHÔNG mã hoá) trong thư mục người dùng -
# tiện cho máy cá nhân/nội bộ, không nên dùng nếu máy có nhiều người chung tài khoản.
# ----------------------------------------------------------------------------
CONFIG_DIR = os.path.join(os.environ.get("APPDATA") or os.path.expanduser("~"), "VATM_DocAI_Pro")
CONFIG_FILE = os.path.join(CONFIG_DIR, "config.json")


def load_config():
    try:
        with open(CONFIG_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}


def save_config(data):
    os.makedirs(CONFIG_DIR, exist_ok=True)
    with open(CONFIG_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False)


def delete_config_key(key):
    data = load_config()
    if key in data:
        del data[key]
        save_config(data)


# ----------------------------------------------------------------------------
# GIAO DIỆN
# ----------------------------------------------------------------------------
class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("VATM DocAI Pro — Chuyển đổi văn bản PDF sang Word")
        self.geometry("1200x720")
        self.minsize(900, 540)

        self.pdf_path = tk.StringVar()
        self.page_mode = tk.StringVar(value="all")
        self.page_range_text = tk.StringVar()
        self.use_ai = tk.BooleanVar(value=False)
        self.api_key = tk.StringVar()
        self.total_pages = 0
        self.last_pdf_path = None
        self.last_render_items = None
        self.last_page_orientation = {}

        cfg = load_config()
        if cfg.get("api_key"):
            self.api_key.set(cfg["api_key"])
        if cfg.get("use_ai"):
            self.use_ai.set(True)

        self._build_ui()
        self._toggle_page_entry()
        self._toggle_api_entry()

        self.log_queue = queue.Queue()
        self.after(200, self._poll_queue)

    def _build_ui(self):
        pad = {"padx": 10, "pady": 6}

        paned = ttk.PanedWindow(self, orient="horizontal")
        paned.pack(fill="both", expand=True)
        left = ttk.Frame(paned)
        right = ttk.Frame(paned)
        paned.add(left, weight=1)
        paned.add(right, weight=1)

        # ---------------- BÊN TRÁI: điều khiển ----------------
        frm_file = ttk.LabelFrame(left, text="1. Chọn file PDF")
        frm_file.pack(fill="x", **pad)
        ttk.Entry(frm_file, textvariable=self.pdf_path).pack(side="left", padx=8, pady=8, fill="x", expand=True)
        ttk.Button(frm_file, text="Chọn file...", command=self.browse_pdf).pack(side="left", padx=8, pady=8)

        frm_pages = ttk.LabelFrame(left, text="2. Chọn trang cần chuyển đổi")
        frm_pages.pack(fill="x", **pad)
        ttk.Radiobutton(frm_pages, text="Tất cả các trang", variable=self.page_mode, value="all",
                         command=self._toggle_page_entry).pack(anchor="w", padx=8, pady=2)
        row = ttk.Frame(frm_pages)
        row.pack(anchor="w", fill="x", padx=8, pady=2)
        ttk.Radiobutton(row, text="Chọn trang:", variable=self.page_mode, value="range",
                         command=self._toggle_page_entry).pack(side="left")
        self.entry_pages = ttk.Entry(row, textvariable=self.page_range_text, width=30, state="disabled")
        self.entry_pages.pack(side="left", padx=8)
        ttk.Label(row, text="ví dụ: 1-3,5,8-10").pack(side="left")
        self.lbl_total_pages = ttk.Label(frm_pages, text="(chưa chọn file)")
        self.lbl_total_pages.pack(anchor="w", padx=8, pady=2)

        frm_ai = ttk.LabelFrame(left, text="3. (Tuỳ chọn) Dùng AI chuẩn hoá trình bày")
        frm_ai.pack(fill="x", **pad)
        ttk.Checkbutton(frm_ai, text="Bật AI (Claude API) để tự chỉnh cấp tiêu đề và sửa lỗi OCR còn sót",
                         variable=self.use_ai, command=self._toggle_api_entry).pack(anchor="w", padx=8, pady=2)
        row2 = ttk.Frame(frm_ai)
        row2.pack(anchor="w", fill="x", padx=8, pady=2)
        ttk.Label(row2, text="Anthropic API key:").pack(side="left")
        self.entry_api = ttk.Entry(row2, textvariable=self.api_key, width=34, show="*", state="disabled")
        self.entry_api.pack(side="left", padx=8)
        ttk.Button(row2, text="Lưu API key", command=self.save_api_key).pack(side="left", padx=4)
        ttk.Button(row2, text="Xoá API key", command=self.delete_api_key).pack(side="left", padx=4)

        frm_run = ttk.Frame(left)
        frm_run.pack(fill="x", **pad)
        self.btn_start = ttk.Button(frm_run, text="Bắt đầu chuyển đổi", command=self.start_processing)
        self.btn_start.pack(side="left")
        self.progress = ttk.Progressbar(frm_run, mode="indeterminate")
        self.progress.pack(side="left", padx=10, fill="x", expand=True)

        frm_log = ttk.LabelFrame(left, text="Tiến trình")
        frm_log.pack(fill="both", expand=True, **pad)
        self.txt_log = tk.Text(frm_log, height=14, state="disabled", wrap="word")
        self.txt_log.pack(fill="both", expand=True, padx=8, pady=8)

        # ---------------- BÊN PHẢI: xem trước + xuất file ----------------
        frm_preview_header = ttk.Frame(right)
        frm_preview_header.pack(fill="x", padx=8, pady=(8, 0))
        ttk.Label(frm_preview_header, text="Xem trước kết quả", font=("", 10, "bold")).pack(side="left")
        self.btn_export = ttk.Button(frm_preview_header, text="Xuất file Word...",
                                      command=self.export_word, state="disabled")
        self.btn_export.pack(side="right", padx=4)
        self.btn_enlarge = ttk.Button(frm_preview_header, text="Phóng to",
                                       command=self.open_enlarged_preview, state="disabled")
        self.btn_enlarge.pack(side="right", padx=4)

        frm_preview = ttk.Frame(right)
        frm_preview.pack(fill="both", expand=True, padx=8, pady=8)
        yscroll = ttk.Scrollbar(frm_preview, orient="vertical")
        self.txt_preview = tk.Text(frm_preview, state="disabled", wrap="word", yscrollcommand=yscroll.set)
        yscroll.config(command=self.txt_preview.yview)
        yscroll.pack(side="right", fill="y")
        self.txt_preview.pack(side="left", fill="both", expand=True)

    def _toggle_page_entry(self):
        self.entry_pages.config(state="normal" if self.page_mode.get() == "range" else "disabled")

    def _toggle_api_entry(self):
        self.entry_api.config(state="normal" if self.use_ai.get() else "disabled")

    def save_api_key(self):
        key = self.api_key.get().strip()
        if not key:
            messagebox.showwarning("Thiếu API key", "Chưa nhập API key để lưu.")
            return
        cfg = load_config()
        cfg["api_key"] = key
        cfg["use_ai"] = self.use_ai.get()
        save_config(cfg)
        messagebox.showinfo("Đã lưu", f"Đã lưu API key vào:\n{CONFIG_FILE}\n"
                                       f"(lưu ý: file lưu dạng thường, không mã hoá)")

    def delete_api_key(self):
        delete_config_key("api_key")
        delete_config_key("use_ai")
        self.api_key.set("")
        messagebox.showinfo("Đã xoá", "Đã xoá API key đã lưu trên máy.")

    def browse_pdf(self):
        path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
        if not path:
            return
        self.pdf_path.set(path)
        try:
            reader = PdfReader(path)
            self.total_pages = len(reader.pages)
            self.lbl_total_pages.config(text=f"File có {self.total_pages} trang.")
        except Exception as ex:
            messagebox.showerror("Lỗi", f"Không đọc được file PDF: {ex}")

    def log(self, msg):
        self.txt_log.config(state="normal")
        self.txt_log.insert("end", msg + "\n")
        self.txt_log.see("end")
        self.txt_log.config(state="disabled")

    def start_processing(self):
        pdf_path = self.pdf_path.get().strip()
        if not pdf_path or not os.path.isfile(pdf_path):
            messagebox.showwarning("Thiếu thông tin", "Hãy chọn file PDF hợp lệ trước.")
            return

        if self.total_pages == 0:
            messagebox.showwarning("Thiếu thông tin", "Không đọc được số trang của file này.")
            return

        if self.page_mode.get() == "all":
            page_list = list(range(1, self.total_pages + 1))
        else:
            try:
                page_list = parse_page_selection(self.page_range_text.get(), self.total_pages)
            except Exception as ex:
                messagebox.showwarning("Sai định dạng", f"{ex}\nVí dụ đúng: 1-3,5,8-10")
                return

        use_ai = self.use_ai.get()
        api_key = self.api_key.get().strip()
        if use_ai and not api_key:
            messagebox.showwarning("Thiếu API key", "Đã bật AI nhưng chưa nhập Anthropic API key.")
            return

        output_path = os.path.splitext(pdf_path)[0] + "_ketqua.docx"
        self.last_pdf_path = pdf_path
        self.last_render_items = None

        self.btn_start.config(state="disabled")
        self.btn_export.config(state="disabled")
        self.btn_enlarge.config(state="disabled")
        self.progress.start(12)
        self.log(f"Bắt đầu xử lý {len(page_list)} trang: {page_list}")

        self.txt_preview.config(state="normal")
        self.txt_preview.delete("1.0", "end")
        self.txt_preview.config(state="disabled")

        thread = threading.Thread(
            target=process_pdf_safe,
            args=(pdf_path, page_list, output_path, use_ai, api_key, self.log_queue),
            daemon=True,
        )
        thread.start()

    def export_word(self):
        if not self.last_render_items:
            messagebox.showwarning("Chưa có dữ liệu", "Chưa có kết quả để xuất - hãy chuyển đổi trước.")
            return
        default_name = "ket_qua.docx"
        if self.last_pdf_path:
            default_name = os.path.splitext(os.path.basename(self.last_pdf_path))[0] + "_ketqua.docx"
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word document", "*.docx")],
            initialfile=default_name,
        )
        if not path:
            return
        try:
            render_docx(self.last_render_items, path, page_orientation=self.last_page_orientation)
            self.log(f"Đã xuất file Word: {path}")
            messagebox.showinfo("Xong", f"Đã lưu file:\n{path}")
        except Exception as ex:
            messagebox.showerror("Lỗi xuất file", str(ex))

    def open_enlarged_preview(self):
        if not self.last_render_items:
            return
        win = tk.Toplevel(self)
        win.title("Xem trước (phóng to)")
        win.geometry("1000x820")
        yscroll = ttk.Scrollbar(win, orient="vertical")
        txt = tk.Text(win, wrap="word", yscrollcommand=yscroll.set, font=("Times New Roman", 13))
        yscroll.config(command=txt.yview)
        yscroll.pack(side="right", fill="y")
        txt.pack(side="left", fill="both", expand=True)
        render_preview_into_widget(txt, self.last_render_items)

    def _poll_queue(self):
        try:
            while True:
                item = self.log_queue.get_nowait()
                kind = item[0]
                if kind == "status":
                    self.log(item[1])
                elif kind == "error":
                    self.log("LỖI:\n" + item[1])
                    self.progress.stop()
                    self.btn_start.config(state="normal")
                    messagebox.showerror("Lỗi xử lý", "Có lỗi xảy ra, xem chi tiết trong khung Tiến trình.")
                elif kind == "preview_ready":
                    render_items, page_orientation, sidebar_notes = item[1], item[2], item[3]
                    self.last_render_items = render_items
                    self.last_page_orientation = page_orientation
                    render_preview_into_widget(self.txt_preview, render_items)
                    self.log("\nĐÃ XONG - xem bản xem trước bên phải, bấm \"Xuất file Word...\" để lưu file.")
                    for n in sidebar_notes:
                        self.log("  (lề/tiêu đề bị tách riêng) " + n)
                    self.progress.stop()
                    self.btn_start.config(state="normal")
                    self.btn_export.config(state="normal")
                    self.btn_enlarge.config(state="normal")
        except queue.Empty:
            pass
        self.after(200, self._poll_queue)


if __name__ == "__main__":
    app = App()
    app.mainloop()
