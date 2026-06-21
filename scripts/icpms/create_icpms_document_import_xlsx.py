import sys
import os
import re
from datetime import datetime
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.worksheet.datavalidation import DataValidation
from openpyxl.utils import get_column_letter

def clean_text(text):
    if not text:
        return ""
    return text.strip().replace('\n', ' ')

def main():
    docx_path = os.path.join("docs", "2026.03.19 Tai_lieu_ICAO_IMS_VATM.docx")
    out_path = os.path.join("data", "import", "icpms_document_catalog_import.xlsx")
    
    if not os.path.exists(docx_path):
        print(f"File not found: {docx_path}")
        return

    doc = Document(docx_path)
    
    documents = []
    current_section = ""
    
    # Define dictionaries
    dict_document_type = [
        "ICAO_ANNEX", "ICAO_DOC", "ICAO_CIRCULAR", "ICAO_APAC", "CANSO_GUIDANCE", 
        "ISO_STANDARD", "EASA_EU", "EUROCONTROL", "EUROCAE_RTCA", "VATM_INTERNAL",
        "LAW", "DECREE", "CIRCULAR_VN", "DECISION", "DIRECTIVE", "OFFICIAL_LETTER",
        "NATIONAL_STANDARD", "TECHNICAL_REGULATION", "INTERNAL_REGULATION", "PROCEDURE",
        "GUIDANCE", "FORM", "SAFETY_DOCUMENT", "QUALITY_DOCUMENT", "SECURITY_DOCUMENT",
        "COMPLIANCE_DOCUMENT", "OTHER"
    ]
    dict_document_group = [
        "ICAO", "ICAO_APAC", "CANSO", "ISO", "EASA_EU", "EUROCONTROL", "EUROCAE_RTCA",
        "VIETNAM_LEGAL", "CAAV", "VATM", "INTERNAL", "OTHER"
    ]
    dict_language = ["English", "Vietnamese"]
    dict_classification = ["Public", "Internal", "Restricted", "Confidential"]
    dict_app_vatm = ["YES", "NO", "REVIEW"]
    dict_priority = ["HIGH", "MEDIUM", "LOW"]
    dict_status = ["DRAFT", "ACTIVE", "UNDER_REVIEW", "SUPERSEDED", "ARCHIVED"]

    # Parse DOCX
    for element in doc.element.body:
        if element.tag.endswith('p'):
            # paragraph
            text = ""
            for node in element.iter():
                if node.tag.endswith('t') and node.text:
                    text += node.text
            text = text.upper()
            if "APAC" in text:
                current_section = "ICAO_APAC"
            elif "CANSO" in text:
                current_section = "CANSO"
            elif "ISO" in text:
                current_section = "ISO"
            elif "EASA" in text or "EU" in text:
                current_section = "EASA_EU"
            elif "EUROCONTROL" in text:
                current_section = "EUROCONTROL"
            elif "EUROCAE" in text or "RTCA" in text:
                current_section = "EUROCAE_RTCA"
                
        elif element.tag.endswith('tbl'):
            # This logic depends on python-docx handling elements. 
            # A safer way is to iterate doc.blocks (or just iterate through tables and hope sections are determinable)
            pass

    # Actually, python-docx iterates paragraphs and tables separately using doc.paragraphs and doc.tables, 
    # but their order is lost unless we iterate element.body.
    # We will do it properly:
    current_section = "ICAO"
    for block in doc.element.body:
        if block.tag.endswith('p'):
            p_text = "".join(node.text for node in block.iter() if node.tag.endswith('t') and node.text).upper()
            if "APAC" in p_text: current_section = "ICAO_APAC"
            elif "CANSO" in p_text: current_section = "CANSO"
            elif "ISO" in p_text: current_section = "ISO"
            elif "EASA" in p_text or "EU " in p_text: current_section = "EASA_EU"
            elif "EUROCONTROL" in p_text: current_section = "EUROCONTROL"
            elif "EUROCAE" in p_text or "RTCA" in p_text: current_section = "EUROCAE_RTCA"
        elif block.tag.endswith('tbl'):
            # Find the corresponding Table object
            table = None
            for tbl in doc.tables:
                if tbl._element == block:
                    table = tbl
                    break
            
            if table:
                headers = [clean_text(cell.text).lower() for cell in table.rows[0].cells]
                
                # Check if it looks like a document table
                code_idx, title_idx, domain_idx, page_idx = -1, -1, -1, -1
                for i, h in enumerate(headers):
                    if "mã" in h or "code" in h: code_idx = i
                    elif "tên" in h or "title" in h or "document" in h: title_idx = i
                    elif "lĩnh vực" in h or "domain" in h: domain_idx = i
                    elif "trang" in h or "page" in h or "pp" in h: page_idx = i
                
                # Default indices if headers are merged/missing but standard format
                if code_idx == -1 and len(headers) >= 4:
                    code_idx, title_idx = 2, 3
                    
                if code_idx != -1 and title_idx != -1:
                    for row in table.rows[1:]:
                        cells = row.cells
                        if len(cells) > max(code_idx, title_idx):
                            code = clean_text(cells[code_idx].text)
                            title = clean_text(cells[title_idx].text)
                            
                            # Skip empty or header repeats
                            if not code or code.lower() == "mã tài liệu" or code.lower() == "mã": continue
                            
                            domain = clean_text(cells[domain_idx].text) if domain_idx != -1 and len(cells) > domain_idx else ""
                            pages_text = clean_text(cells[page_idx].text) if page_idx != -1 and len(cells) > page_idx else ""
                            
                            page_count = ""
                            if pages_text:
                                nums = re.findall(r'\d+', pages_text)
                                if nums: page_count = nums[0]

                            doc_type = "OTHER"
                            doc_group = "OTHER"
                            org = ""

                            if code.startswith("Annex") or code.startswith("ANNEX"):
                                doc_type = "ICAO_ANNEX"
                                doc_group = "ICAO"
                                org = "ICAO"
                            elif code.startswith("Doc ") or code.startswith("DOC "):
                                doc_type = "ICAO_DOC"
                                doc_group = "ICAO"
                                org = "ICAO"
                            elif code.startswith("Cir ") or code.startswith("CIR "):
                                doc_type = "ICAO_CIRCULAR"
                                doc_group = "ICAO"
                                org = "ICAO"
                            elif code.startswith("ISO") or current_section == "ISO":
                                doc_type = "ISO_STANDARD"
                                doc_group = "ISO"
                                org = "ISO"
                            elif current_section == "ICAO_APAC" or "APAC" in code:
                                doc_type = "ICAO_APAC"
                                doc_group = "ICAO_APAC"
                                org = "ICAO APAC"
                            elif current_section == "CANSO":
                                doc_type = "CANSO_GUIDANCE"
                                doc_group = "CANSO"
                                org = "CANSO"
                            elif current_section == "EASA_EU":
                                doc_type = "EASA_EU"
                                doc_group = "EASA_EU"
                                org = "EASA/EU"
                            elif current_section == "EUROCONTROL":
                                doc_type = "EUROCONTROL"
                                doc_group = "EUROCONTROL"
                                org = "EUROCONTROL"
                            elif current_section == "EUROCAE_RTCA":
                                doc_type = "EUROCAE_RTCA"
                                doc_group = "EUROCAE_RTCA"
                                org = "EUROCAE/RTCA"

                            documents.append({
                                "code": code,
                                "title": title,
                                "type": doc_type,
                                "group": doc_group,
                                "org": org,
                                "domain": domain,
                                "pages": page_count,
                                "lang": "English",
                                "class": "Public",
                                "app": "REVIEW",
                                "priority": "MEDIUM",
                                "status": "ACTIVE",
                                "notes": "Trích xuất từ danh mục tài liệu ICAO IMS VATM"
                            })

    extracted_count = len(documents)

    # Append VN templates
    vn_templates = [
        {"code": "VN-LEGAL-001", "title": "Luật Hàng không dân dụng Việt Nam", "type": "LAW", "group": "VIETNAM_LEGAL", "org": "Quốc hội", "domain": "", "pages": "", "lang": "Vietnamese", "class": "Public", "app": "REVIEW", "priority": "HIGH", "status": "ACTIVE", "notes": ""},
        {"code": "ND-32-2016-ND-CP", "title": "Nghị định 32/2016/NĐ-CP", "type": "DECREE", "group": "VIETNAM_LEGAL", "org": "Chính phủ", "domain": "", "pages": "", "lang": "Vietnamese", "class": "Public", "app": "REVIEW", "priority": "HIGH", "status": "ACTIVE", "notes": ""},
        {"code": "CAAV-SAFETY-GUIDANCE-SAMPLE", "title": "Văn bản hướng dẫn của Cục Hàng không Việt Nam về an toàn hàng không", "type": "GUIDANCE", "group": "CAAV", "org": "Cục Hàng không Việt Nam", "domain": "", "pages": "", "lang": "Vietnamese", "class": "Public", "app": "REVIEW", "priority": "HIGH", "status": "ACTIVE", "notes": ""},
        {"code": "VATM-SMS-MANUAL", "title": "Sổ tay quản lý an toàn VATM", "type": "SAFETY_DOCUMENT", "group": "VATM", "org": "VATM", "domain": "", "pages": "", "lang": "Vietnamese", "class": "Internal", "app": "YES", "priority": "HIGH", "status": "DRAFT", "notes": ""},
        {"code": "VATM-COMPLIANCE-PROCEDURE", "title": "Quy trình quản lý tuân thủ VATM", "type": "PROCEDURE", "group": "VATM", "org": "VATM", "domain": "", "pages": "", "lang": "Vietnamese", "class": "Internal", "app": "YES", "priority": "HIGH", "status": "ACTIVE", "notes": ""},
        {"code": "VATM-EVIDENCE-GUIDE", "title": "Hướng dẫn quản lý bằng chứng tuân thủ", "type": "GUIDANCE", "group": "VATM", "org": "VATM", "domain": "", "pages": "", "lang": "Vietnamese", "class": "Internal", "app": "YES", "priority": "HIGH", "status": "ACTIVE", "notes": ""},
        {"code": "VATM-DOCUMENT-CONTROL-PROCEDURE", "title": "Quy trình kiểm soát tài liệu và hồ sơ VATM", "type": "PROCEDURE", "group": "VATM", "org": "VATM", "domain": "", "pages": "", "lang": "Vietnamese", "class": "Internal", "app": "YES", "priority": "HIGH", "status": "ACTIVE", "notes": ""}
    ]
    
    documents.extend(vn_templates)

    # Validation
    errors = 0
    validation_rows = []
    
    for idx, d in enumerate(documents):
        row_num = idx + 2
        
        # Priority rules for VN docs
        if d["lang"] == "Vietnamese":
            title_upper = d["title"].upper()
            if any(k in title_upper for k in ["SAFETY", "SMS", "COMPLIANCE", "ATS", "CNS", "MET", "AIM", "SAR", "ATFM", "AN TOÀN", "TUÂN THỦ"]):
                d["priority"] = "HIGH"

        # Validations
        if not d["code"]:
            errors += 1
            validation_rows.append([row_num, "document_code", "Trống", "Mã tài liệu không được trống"])
        if not d["title"]:
            errors += 1
            validation_rows.append([row_num, "document_title", "Trống", "Tên tài liệu không được trống"])
        if d["type"] not in dict_document_type:
            errors += 1
            validation_rows.append([row_num, "document_type", d["type"], "Không thuộc danh mục hợp lệ"])
        if d["status"] not in dict_status:
            errors += 1
            validation_rows.append([row_num, "status", d["status"], "Không thuộc danh mục hợp lệ"])
        if d["pages"] and not str(d["pages"]).isdigit():
            errors += 1
            validation_rows.append([row_num, "page_count", d["pages"], "Phải là số nguyên"])
        if d["app"] == "REVIEW":
            validation_rows.append([row_num, "applicable_to_vatm", "REVIEW", "Cần VATM rà soát mức độ áp dụng"])

    # Create Excel
    wb = Workbook()
    
    # 1. documents sheet
    ws_doc = wb.active
    ws_doc.title = "documents"
    headers = [
        "stt", "document_code", "document_title", "document_type", "document_group", 
        "source_organization", "main_domain", "page_count", "issued_date", 
        "effective_date", "language", "classification", "applicable_to_vatm", 
        "priority", "status", "notes"
    ]
    ws_doc.append(headers)
    
    # Header format
    header_fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
    header_font = Font(bold=True)
    for col_idx, cell in enumerate(ws_doc[1], 1):
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
    
    # Freeze
    ws_doc.freeze_panes = "A2"
    
    # Data
    for idx, d in enumerate(documents):
        ws_doc.append([
            idx + 1, d["code"], d["title"], d["type"], d["group"], d["org"], d["domain"], 
            d["pages"], "", "", d["lang"], d["class"], d["app"], d["priority"], d["status"], d["notes"]
        ])
    
    # Validation Dropdowns
    dv_type = DataValidation(type="list", formula1="=dictionaries!$A$2:$A$100", allow_blank=True)
    dv_group = DataValidation(type="list", formula1="=dictionaries!$B$2:$B$100", allow_blank=True)
    dv_lang = DataValidation(type="list", formula1="=dictionaries!$C$2:$C$100", allow_blank=True)
    dv_class = DataValidation(type="list", formula1="=dictionaries!$D$2:$D$100", allow_blank=True)
    dv_app = DataValidation(type="list", formula1="=dictionaries!$E$2:$E$100", allow_blank=True)
    dv_priority = DataValidation(type="list", formula1="=dictionaries!$F$2:$F$100", allow_blank=True)
    dv_status = DataValidation(type="list", formula1="=dictionaries!$G$2:$G$100", allow_blank=True)

    ws_doc.add_data_validation(dv_type)
    ws_doc.add_data_validation(dv_group)
    ws_doc.add_data_validation(dv_lang)
    ws_doc.add_data_validation(dv_class)
    ws_doc.add_data_validation(dv_app)
    ws_doc.add_data_validation(dv_priority)
    ws_doc.add_data_validation(dv_status)

    dv_type.add(f"D2:D1000")
    dv_group.add(f"E2:E1000")
    dv_lang.add(f"K2:K1000")
    dv_class.add(f"L2:L1000")
    dv_app.add(f"M2:M1000")
    dv_priority.add(f"N2:N1000")
    dv_status.add(f"O2:O1000")

    # Layout
    ws_doc.auto_filter.ref = ws_doc.dimensions
    col_widths = {"A": 5, "B": 25, "C": 50, "D": 20, "E": 20, "F": 20, "G": 15, "H": 10, "K": 15, "L": 15, "M": 15, "N": 15, "O": 15, "P": 30}
    for col, width in col_widths.items():
        ws_doc.column_dimensions[col].width = width

    # 2. dictionaries sheet
    ws_dict = wb.create_sheet("dictionaries")
    dict_headers = ["document_type", "document_group", "language", "classification", "applicable_to_vatm", "priority", "status"]
    ws_dict.append(dict_headers)
    
    max_len = max(len(dict_document_type), len(dict_document_group), len(dict_language), len(dict_classification), len(dict_app_vatm), len(dict_priority), len(dict_status))
    
    for i in range(max_len):
        row = []
        row.append(dict_document_type[i] if i < len(dict_document_type) else "")
        row.append(dict_document_group[i] if i < len(dict_document_group) else "")
        row.append(dict_language[i] if i < len(dict_language) else "")
        row.append(dict_classification[i] if i < len(dict_classification) else "")
        row.append(dict_app_vatm[i] if i < len(dict_app_vatm) else "")
        row.append(dict_priority[i] if i < len(dict_priority) else "")
        row.append(dict_status[i] if i < len(dict_status) else "")
        ws_dict.append(row)
        
    for col_idx, cell in enumerate(ws_dict[1], 1):
        cell.font = header_font

    # 3. import_notes sheet
    ws_notes = wb.create_sheet("import_notes")
    ws_notes.append(["Metric", "Value"])
    ws_notes.append(["Tên file nguồn", "2026.03.19 Tai_lieu_ICAO_IMS_VATM.docx"])
    ws_notes.append(["Ngày tạo file", datetime.now().strftime("%Y-%m-%d %H:%M:%S")])
    ws_notes.append(["Tổng số dòng trích xuất", extracted_count])
    ws_notes.append(["Số dòng tiếng Anh/quốc tế", extracted_count])
    ws_notes.append(["Số dòng tiếng Việt/VATM mẫu", len(vn_templates)])
    ws_notes.append(["Số dòng cần rà soát", len([1 for d in documents if d["app"] == "REVIEW"])])
    ws_notes.append(["Ghi chú", "File này chỉ là metadata danh mục tài liệu, không phải file bóc tách requirements."])
    ws_notes.column_dimensions["A"].width = 30
    ws_notes.column_dimensions["B"].width = 50

    # 4. validation_report sheet
    ws_val = wb.create_sheet("validation_report")
    ws_val.append(["Row", "Column", "Value", "Issue"])
    for r in validation_rows:
        ws_val.append(r)
        
    for col_idx, cell in enumerate(ws_val[1], 1):
        cell.font = header_font
    ws_val.column_dimensions["A"].width = 10
    ws_val.column_dimensions["B"].width = 25
    ws_val.column_dimensions["C"].width = 25
    ws_val.column_dimensions["D"].width = 50

    wb.save(out_path)
    
    print(f"XLSX saved at {out_path}")
    print(f"Total documents: {len(documents)}")
    print(f"English/Intl: {extracted_count}")
    print(f"Vietnamese/VATM: {len(vn_templates)}")
    print(f"Need Review: {len([1 for d in documents if d['app'] == 'REVIEW'])}")
    print(f"Validation Errors: {errors}")

if __name__ == "__main__":
    main()
